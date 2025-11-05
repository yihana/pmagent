# server/workflow/agents/scope_agent/output/
from pathlib import Path
from typing import Dict, Any, List

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ProjectCharterGenerator:
    """프로젝트 헌장 Word 문서 생성기"""
    
    @staticmethod
    def generate(
        project_name: str,
        requirements: List[Dict[str, Any]],
        wbs_data: Dict[str, Any],
        output_path: Path
    ) -> str:
        """
        프로젝트 헌장 Word 문서 생성
        
        Args:
            project_name: 프로젝트명
            requirements: 요구사항 리스트
            wbs_data: WBS 데이터
            output_path: 저장할 파일 경로
            
        Returns:
            str: 생성된 파일의 절대 경로
        """
        if not DOCX_AVAILABLE:
            # python-docx 미설치 시 텍스트 파일로 대체
            output_path = output_path.with_suffix('.txt')
            output_path.write_text(
                f"프로젝트 헌장\n\n프로젝트명: {project_name}\n\n"
                "[python-docx 미설치로 인해 텍스트 파일로 생성됨]",
                encoding="utf-8"
            )
            return str(output_path.resolve())
        
        doc = Document()
        
        # 제목
        title = doc.add_heading("프로젝트 헌장", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        
        doc.add_paragraph()  # 공백
        
        # 표 생성 (Image 9 양식)
        sections = [
            ("항목", "세부설명", "비고"),
            ("프로젝트 기본정보", ProjectCharterGenerator._get_basic_info(project_name), ""),
            ("프로젝트 개요 및 추진배경", ProjectCharterGenerator._get_overview(), ""),
            ("프로젝트 목적 및 목표(Outcome Benefits)", ProjectCharterGenerator._get_objectives(), ""),
            ("프로젝트 성공기준", ProjectCharterGenerator._get_success_criteria(), ""),
            ("프로젝트 성공요인", ProjectCharterGenerator._get_success_factors(), ""),
            ("성공 추진 전략", ProjectCharterGenerator._get_strategy(), ""),
            ("추진방법론 및 표준", ProjectCharterGenerator._get_methodology(), ""),
            ("주요 변경 및 산출물(인도물)", ProjectCharterGenerator._get_deliverables(wbs_data), ""),
            ("주요 이해관계자 식별", ProjectCharterGenerator._get_stakeholders(), ""),
            ("PM 권한 및 역할", ProjectCharterGenerator._get_pm_authority(), ""),
            ("가정 및 전제조건", ProjectCharterGenerator._get_assumptions(), ""),
            ("주요 위험 및 대응(작수사)", ProjectCharterGenerator._get_risks(), ""),
            ("원가내역", ProjectCharterGenerator._get_cost(), ""),
            ("주요 마일스톤", ProjectCharterGenerator._get_milestones(), ""),
            ("조직 및 역할", ProjectCharterGenerator._get_organization(), ""),
            ("스폰서 및 보고라인", ProjectCharterGenerator._get_sponsor(), ""),
        ]
        
        table = doc.add_table(rows=len(sections), cols=3)
        table.style = 'Light Grid Accent 1'
        
        # 표 데이터 입력
        for i, (col1, col2, col3) in enumerate(sections):
            row_cells = table.rows[i].cells
            row_cells[0].text = col1
            row_cells[1].text = col2
            row_cells[2].text = col3
            
            # 헤더 스타일
            if i == 0:
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
            else:
                # 항목명 굵게
                for paragraph in row_cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
        
        # 열 너비 조정
        for row in table.rows:
            row.cells[0].width = Inches(2.0)
            row.cells[1].width = Inches(4.5)
            row.cells[2].width = Inches(1.0)
        
        # 파일 저장
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        return str(output_path.resolve())
    
    @staticmethod
    def _get_basic_info(project_name: str) -> str:
        return f"프로젝트명: {project_name}\n기간: 2022.06 ~ 2023.03\n관리자: 3조원"
    
    @staticmethod
    def _get_overview() -> str:
        return "무원별 역량향상을 위한 학습 플랫폼 구축"
    
    @staticmethod
    def _get_objectives() -> str:
        return "• 기본기능 중심 안정적 시스템 오픈\n• 전략기능 기반 시스템 고도화\n• 외부 Open 추진 및 DT기반 시스템 혁신"
    
    @staticmethod
    def _get_success_criteria() -> str:
        return "• 인수 테스트 통과\n• 사용자 만족도 80% 이상\n• 시스템 가용성 99.9% 이상"
    
    @staticmethod
    def _get_success_factors() -> str:
        return "• 명확한 요구사항 정의\n• 효과적인 이해관계자 관리\n• 적절한 리소스 배분"
    
    @staticmethod
    def _get_strategy() -> str:
        return "• Agile 방법론 적용\n• 단계적 릴리스\n• 지속적 통합/배포(CI/CD)"
    
    @staticmethod
    def _get_methodology() -> str:
        return "• PMP 표준 프로세스\n• Agile/Scrum\n• DevOps"
    
    @staticmethod
    def _get_deliverables(wbs_data: Dict) -> str:
        deliverables = [
            "1. 착수보고서",
            "2. 요구사항정의서",
            "3. WBS",
            "4. 테스트 결과서",
            "5. 사용자 매뉴얼"
        ]
        return "\n".join(deliverables)
    
    @staticmethod
    def _get_stakeholders() -> str:
        return "• Sponsor: 경영진\n• PM: 프로젝트 관리자\n• 개발팀: 개발자\n• 사용자: 최종 사용자"
    
    @staticmethod
    def _get_pm_authority() -> str:
        return "• 예산 승인 권한\n• 리소스 배분 권한\n• 일정 조정 권한\n• 변경 관리 권한"
    
    @staticmethod
    def _get_assumptions() -> str:
        return "• 그룹사 입직원 대상 서비스\n• 기존 인프라 활용 가능\n• 관련 부서 협조 가능"
    
    @staticmethod
    def _get_risks() -> str:
        return "• 요구사항 변경 리스크 (대응: 변경관리 프로세스 수립)\n• 일정 지연 리스크 (대응: 버퍼 확보)\n• 기술 리스크 (대응: POC 수행)"
    
    @staticmethod
    def _get_cost() -> str:
        return "• 인건비: 70%\n• 외주비: 20%\n• 기타: 10%\n• 총 예산: TBD"
    
    @staticmethod
    def _get_milestones() -> str:
        return "• 착수: 2022.06\n• 설계완료: 2022.09\n• 개발완료: 2023.01\n• 오픈: 2023.03"
    
    @staticmethod
    def _get_organization() -> str:
        return "• PM: 프로젝트 총괄\n• 개발팀: 시스템 개발\n• QA팀: 품질 관리\n• 인프라팀: 인프라 구축"
    
    @staticmethod
    def _get_sponsor() -> str:
        return "• Sponsor: CEO\n• 보고 주기: 주간\n• 보고 방식: 서면/회의"