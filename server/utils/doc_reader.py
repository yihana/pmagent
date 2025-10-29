# server/utils/doc_reader.py
from __future__ import annotations
from pathlib import Path
from typing import Optional, Tuple, List, Union
import mimetypes

# --------------------------------------
# Exceptions
# --------------------------------------
class DocReadError(Exception):
    pass

# --------------------------------------
# Path helpers
# --------------------------------------
def resolve_path(path: Union[str, Path], search_roots: Optional[List[Path]] = None) -> Optional[Path]:
    """
    상대/부분 경로를 여러 후보 루트에서 탐색해 실제 파일 경로를 찾는다.
    찾지 못하면 None 반환.
    """
    p = Path(str(path)).expanduser()
    candidates: List[Path] = []
    if p.is_absolute():
        candidates.append(p)

    # 현재 작업 디렉토리 기준 후보들
    candidates += [Path.cwd() / p, p, Path.cwd() / p.name]

    # 제공된 검색 루트들
    for root in (search_roots or []):
        root = Path(root)
        candidates += [root / p, root / p.name]

    for c in candidates:
        try:
            if c.exists() and c.is_file():
                return c.resolve()
        except Exception:
            continue
    return None

# --------------------------------------
# Text readers
# --------------------------------------
def _read_text_utf8(p: Path) -> str:
    try:
        return p.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return p.read_bytes().decode("utf-8", errors="ignore")

def _read_pdf_text(p: Path) -> Tuple[str, str]:
    """
    PDF 텍스트 추출 (다중 폴백). (text, extractor_name) 반환.
    폴백 순서는 PyMuPDF -> pypdf -> pdfplumber.
    """
    # 1) PyMuPDF
    try:
        import fitz  # PyMuPDF
        chunks = []
        with fitz.open(p) as doc:
            for page in doc:
                chunks.append(page.get_text("text") or page.get_text() or "")
        txt = "\n".join(chunks).strip()
        if txt:
            return txt, "PyMuPDF"
    except Exception:
        pass

    # 2) pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(p))
        txt = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        if txt:
            return txt, "pypdf"
    except Exception:
        pass

    # 3) pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(str(p)) as pdf:
            txt = "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
        if txt:
            return txt, "pdfplumber"
    except Exception:
        pass

    return "", "none"

def _read_docx_text(p: Path) -> str:
    """
    DOCX 문서에서:
      - 본문 문단(글머리표 감지)
      - 표(마크다운 테이블로 변환)
      - 1번째 섹션 헤더/푸터
    까지 텍스트화. 내용이 비어있으면 docx2txt fallback 시도.
    """
    try:
        import docx  # pip install python-docx
    except Exception:
        raise DocReadError("DOCX 추출을 위해 `pip install python-docx`가 필요합니다.")

    try:
        d = docx.Document(str(p))
        lines: list[str] = []

        # 1) 본문 문단
        for para in d.paragraphs:
            t = (para.text or "").strip()
            if not t:
                continue
            style_name = str(getattr(getattr(para, "style", None), "name", ""))  # 안전 접근
            if any(k in style_name for k in ("List", "Bulleted", "Numbered")):
                lines.append(f"- {t}")
            else:
                lines.append(t)

        # 2) 표 → 마크다운 테이블
        def cell_text(cell) -> str:
            # 셀 내부 여러 문단/개행 정리
            return "\n".join([(p.text or "").strip() for p in cell.paragraphs if (p.text or "").strip()])

        for ti, table in enumerate(d.tables, start=1):
            rows = list(table.rows)
            if not rows:
                continue
            lines.append("")  # 표 앞 빈 줄
            lines.append(f"### Table {ti}")

            headers = [cell_text(c) or f"col{ci+1}" for ci, c in enumerate(rows[0].cells)]
            sep = ["---" for _ in headers]
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(sep) + " |")
            for r in rows[1:]:
                vals = [cell_text(c) for c in r.cells]
                lines.append("| " + " | ".join(vals) + " |")

        # 3) 첫 섹션 헤더/푸터 (있을 때만)
        try:
            sec0 = d.sections[0]
            if getattr(sec0, "header", None):
                htxt = "\n".join([(p.text or "").strip() for p in sec0.header.paragraphs if (p.text or "").strip()])
                if htxt:
                    lines.insert(0, f"[Header] {htxt}")
            if getattr(sec0, "footer", None):
                ftxt = "\n".join([(p.text or "").strip() for p in sec0.footer.paragraphs if (p.text or "").strip()])
                if ftxt:
                    lines.append(f"[Footer] {ftxt}")
        except Exception:
            # 섹션이 없거나 접근 불가해도 무시
            pass

        text = "\n".join(lines).strip()

        # 내용이 여전히 비어있다면 docx2txt로 보조 시도
        if not text:
            try:
                import docx2txt  # pip install docx2txt (선택)
                text = (docx2txt.process(str(p)) or "").strip()
            except Exception:
                pass

        if not text:
            raise DocReadError("DOCX에서 텍스트를 추출하지 못했습니다. (문단/표가 비어있거나 특수 구조)")

        return text

    except DocReadError:
        raise
    except Exception as e:
        raise DocReadError(f"DOCX 추출 오류: {e}")


def read_text_from_path(path_str: str) -> str:
    """단일 파일에서 텍스트 추출 (PDF, DOCX, TXT 등 지원)"""
    p = Path(path_str)
    if not p.exists() or not p.is_file():
        raise DocReadError(f"파일이 존재하지 않습니다: {path_str}")

    suf = p.suffix.lower()
    
    # 텍스트 파일들
    if suf in {".txt", ".md", ".csv", ".tsv", ".log", ".rst"}:
        return _read_text_utf8(p)

    # PDF
    if suf == ".pdf":
        text, extractor = _read_pdf_text(p)
        if not text:
            raise DocReadError("PDF에서 텍스트를 추출하지 못했습니다.")
        return text

    # DOCX
    if suf == ".docx":
        return _read_docx_text(p)

    # MIME 타입 기반 텍스트 파일
    mime, _ = mimetypes.guess_type(str(p))
    if mime and mime.startswith("text/"):
        return _read_text_utf8(p)

    raise DocReadError(f"지원하지 않는 파일 형식입니다: {p.suffix}")


def read_texts(
    paths: list[str], 
    header: bool = True,
    search_roots: Optional[List[Path]] = None
) -> tuple[str, list[dict]]:
    """
    여러 파일에서 텍스트 추출 후 병합
    
    Args:
        paths: 파일 경로 리스트
        header: 파일명 헤더 추가 여부
        search_roots: 상대 경로 탐색 루트 리스트
    
    Returns:
        (병합된 텍스트, 메타데이터 리스트)
    """
    merged, metas = [], []
    
    for path in paths:
        if not path:
            continue
            
        # 경로 해석 (search_roots 사용)
        resolved = resolve_path(path, search_roots=search_roots)
        if not resolved:
            raise DocReadError(f"파일을 찾을 수 없습니다: {path}")
        
        # 텍스트 추출
        text = read_text_from_path(str(resolved))
        
        # 병합
        if header:
            merged.append(f"\n\n=== FILE: {resolved.name} ===\n\n{text}")
        else:
            merged.append(text)
        
        # 메타데이터
        try:
            size = resolved.stat().st_size
        except Exception:
            size = None
        
        metas.append({
            "path": path,
            "resolved_path": str(resolved),
            "size": size
        })
    
    return ("\n".join(merged)).strip(), metas

    
def ingest_text(
    text: Optional[str],
    documents: Optional[list],
    search_roots: Optional[List[Path]] = None,
    header: bool = False
) -> Tuple[str, Optional[str]]:
    """
    텍스트 또는 문서 파일에서 내용 읽기
    
    우선순위: 직접 text > documents[i].path 파일읽기(찾히는 첫 번째)
    
    Args:
        text: 직접 입력된 텍스트
        documents: 문서 리스트 (dict 또는 객체)
        search_roots: 파일 탐색 루트 경로들
        header: 파일명 헤더 추가 여부
    
    Returns:
        (병합 텍스트, 최초로 읽힌 실제 경로)
    """
    # 1) 직접 텍스트가 있으면 우선 사용
    if text and str(text).strip():
        return str(text), None

    # 2) documents에서 파일 읽기
    docs = documents or []
    for d in docs:
        path = d.get("path") if isinstance(d, dict) else getattr(d, "path", None)
        if not path:
            continue
            
        # 경로 해석
        rp = resolve_path(path, search_roots=search_roots)
        if not rp:
            continue
            
        try:
            if header:
                merged, _ = read_texts([str(rp)], header=True, search_roots=search_roots)
                return merged, str(rp)
            else:
                content = read_text_from_path(str(rp))
                if content.strip():
                    return content, str(rp)
        except DocReadError:
            continue
    
    return "", None


__all__ = [
    "DocReadError",
    "resolve_path",
    "read_text_from_path",
    "read_texts",
    "ingest_text",
]